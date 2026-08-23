"""Text extraction from uploaded campaign briefs.

Three formats, because those are the three a media rep actually forwards: `.pdf` (client
decks and RFPs), `.docx` (briefs written in Word), `.txt`. Both parsers are pure Python
(`pypdf`, `python-docx`) so the upload path works on a laptop with no build toolchain.

**This is the only place uploaded bytes are interpreted**, and the contents are untrusted:
nothing is executed, no network or external entity is touched, and every failure mode
returns a status rather than raising into the request. The extracted text is plain
characters by the time anything else sees it.

The cap is applied *while* accumulating, not afterwards. A 20 MB `.docx` is a zip that can
decompress to far more, so building the whole string and then trimming it is how you turn
an upload limit into a memory problem.

Two limits, deliberately different:
- `MAX_EXTRACT_CHARS` is what gets stored. Generous, because the sidecar is a file.
- `AGENT_EXCERPT_CHARS` is what reaches a model. Small, because SOLUTION.md's rule is that
  a whole document never enters agent context.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Literal

from pydantic import BaseModel, Field

from app.config import get_settings

#: The formats we can actually read. The upload endpoint rejects everything else, rather
#: than accepting a file it would silently contribute nothing from.
SUPPORTED_SUFFIXES = (".pdf", ".docx", ".txt")

#: Ceiling on stored text. ~50 pages of prose; beyond that a "brief" is a different kind
#: of document and the rep should say what matters in the chat.
MAX_EXTRACT_CHARS = 200_000

#: Ceiling on what a single tool call hands the model. ~8 pages.
AGENT_EXCERPT_CHARS = 20_000

ExtractionStatus = Literal["ok", "no_text", "unsupported", "failed"]

# C0/C1 control characters except tab and newline. PDF text layers carry these routinely
# and they break JSON transport and terminal rendering downstream.
_CONTROL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
# Written as escapes, not literals: these are invisible in an editor by definition, and
# a character class of them is unreviewable otherwise.
#   Deleted outright — zero-width joiners, direction marks and the BOM carry no text.
_INVISIBLE = re.compile("[\u200b-\u200f\ufeff]")
#   Folded to a plain space. PDF extraction emits non-breaking and figure spaces
#   heavily; deleting them would join words that were separate on the page.
_ODD_SPACE = re.compile("[\xa0\u2007\u2009\u200a\u202f]")
#   Folded to a newline. U+2028/2029 *are* line breaks, so dropping them would run two
#   lines of a brief together — a budget onto a date, say.
_ODD_BREAK = re.compile("[\u2028\u2029]")
_TRAILING_SPACE = re.compile(r"[ \t]+$", re.MULTILINE)
# Three or more blank lines collapse to one blank line — PDF page breaks produce runs of
# them and they are pure token cost.
_BLANK_RUN = re.compile(r"\n{3,}")


class ExtractedDocument(BaseModel):
    """The result of reading one uploaded file. Never raises; always reports."""

    status: ExtractionStatus
    text: str = ""
    char_count: int = 0
    #: PDFs only. `None` for the formats that have no page concept.
    page_count: int | None = None
    truncated: bool = False
    #: Plain-language reason, safe to show a rep. Empty when `status == "ok"`.
    detail: str = ""


class DocumentSummary(BaseModel):
    """Extraction outcome as recorded on an upload. No document text."""

    extraction_status: ExtractionStatus = "failed"
    char_count: int = 0
    page_count: int | None = None
    truncated: bool = False
    extraction_detail: str = ""
    #: Path to the extracted-text sidecar, relative to the stage root. None if nothing
    #: was extracted.
    text_path: str | None = None
    #: First line or so, for the upload chip in the UI. Not for agent context.
    preview: str = Field(default="", max_length=240)


def clean_text(raw: str) -> str:
    """Normalise extracted text to plain, transport-safe characters."""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = _ODD_BREAK.sub("\n", text)
    text = _CONTROL.sub("", text)
    text = _INVISIBLE.sub("", text)
    text = _ODD_SPACE.sub(" ", text)
    text = _TRAILING_SPACE.sub("", text)
    return _BLANK_RUN.sub("\n\n", text).strip()


def extract_text(path: Path, *, max_chars: int = MAX_EXTRACT_CHARS) -> ExtractedDocument:
    """Read `path` and return its text, or a status explaining why there is none."""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        return ExtractedDocument(
            status="unsupported",
            detail=f"'{suffix}' is not a readable brief format. Supported: "
            f"{', '.join(SUPPORTED_SUFFIXES)}.",
        )
    if not path.exists():
        return ExtractedDocument(status="failed", detail="The staged file is no longer on disk.")

    try:
        if suffix == ".txt":
            return _finish(_read_txt(path, max_chars), max_chars)
        if suffix == ".docx":
            return _finish(_read_docx(path, max_chars), max_chars)
        return _read_pdf(path, max_chars)
    except Exception as exc:  # noqa: BLE001 - an unreadable upload is a status, not a 500
        return ExtractedDocument(
            status="failed",
            detail=f"Could not read this file ({type(exc).__name__}). It may be corrupt.",
        )


def excerpt(text: str, *, limit: int = AGENT_EXCERPT_CHARS) -> tuple[str, bool]:
    """`(text, truncated)` bounded to `limit`, cut on the cleanest nearby boundary.

    Cutting mid-token invites a model to complete the thought, and a fabricated budget
    figure ("the budget is $50,0" becoming "$50,000,000") is the single worst thing this
    system can do. So prefer a paragraph break, then a line break, then at minimum a
    space — only text with no whitespace at all in the last third gets a hard cut.
    """
    if len(text) <= limit:
        return text, False
    window = text[:limit]
    for separator in ("\n\n", "\n", " "):
        cut = window.rfind(separator)
        # Only honour a boundary in the last third, or a document with one early newline
        # would be truncated to almost nothing.
        if cut > limit * 0.66:
            return window[:cut].rstrip(), True
    return window.rstrip(), True


# ------------------------------------------------------------------ per-format readers


def _finish(text: str, max_chars: int) -> ExtractedDocument:
    cleaned = clean_text(text)
    if not cleaned:
        return ExtractedDocument(
            status="no_text",
            detail="The file contains no readable text.",
        )
    return ExtractedDocument(
        status="ok",
        text=cleaned,
        char_count=len(cleaned),
        truncated=len(text) >= max_chars,
    )


def _read_txt(path: Path, max_chars: int) -> str:
    """Decode a text file, tolerating the encodings a Windows desktop produces.

    utf-8 first, then cp1252 — a brief pasted out of Word and saved as .txt is cp1252 far
    more often than it is latin-1, and the two disagree on exactly the characters that
    show up in one (curly quotes, em dashes, the euro sign).
    """
    raw = path.read_bytes()[: max_chars * 4]  # 4 bytes/char worst case for utf-8
    for encoding in ("utf-8", "cp1252"):
        try:
            return raw.decode(encoding)[:max_chars]
        except UnicodeDecodeError:
            continue
    # Last resort: never fail on an encoding, but do not pretend it decoded cleanly.
    return raw.decode("utf-8", errors="replace")[:max_chars]


def _read_docx(path: Path, max_chars: int) -> str:
    """Paragraphs *and* tables.

    Tables are not optional: briefs put the budget, the flight dates and the market list
    in a table roughly as often as in a sentence, and a paragraph-only reader drops
    precisely the fields intake needs.
    """
    import docx  # imported here so an unused format costs nothing at startup

    document = docx.Document(str(path))
    parts: list[str] = []
    total = 0

    def add(chunk: str) -> bool:
        """Append; return False once the cap is reached."""
        nonlocal total
        if not chunk:
            return True
        parts.append(chunk)
        total += len(chunk) + 1
        return total < max_chars

    for paragraph in document.paragraphs:
        if not add(paragraph.text.strip()):
            return "\n".join(parts)

    for table in document.tables:
        for row in table.rows:
            # Tab-separated so a "Budget | $50,000" row survives as one readable line.
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not add("\t".join(c for c in cells if c)):
                return "\n".join(parts)

    return "\n".join(parts)


def _read_pdf(path: Path, max_chars: int) -> ExtractedDocument:
    """Page-by-page text, with the two PDF failures that need distinct messages.

    An encrypted PDF and a scanned one both yield nothing, but the rep's next action is
    completely different — resend it unlocked, versus type the numbers in. Reporting both
    as "no text" would send them down the wrong one.
    """
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
    except (PdfReadError, zipfile.BadZipFile, OSError, ValueError) as exc:
        return ExtractedDocument(
            status="failed",
            detail=f"This PDF could not be opened ({type(exc).__name__}). It may be corrupt.",
        )

    if reader.is_encrypted:
        # Many "encrypted" PDFs carry only an owner password and open with an empty user
        # password, so try before giving up.
        try:
            opened = reader.decrypt("")
        except Exception:  # noqa: BLE001 - unsupported cipher, treated as locked
            opened = 0
        if not opened:
            return ExtractedDocument(
                status="failed",
                page_count=None,
                detail="This PDF is password-protected. Re-upload an unlocked copy.",
            )

    pages = len(reader.pages)
    parts: list[str] = []
    total = 0
    truncated = False
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - one broken page must not lose the rest
            continue
        parts.append(page_text)
        total += len(page_text)
        if total >= max_chars:
            truncated = True
            break

    cleaned = clean_text("\n\n".join(parts))
    if not cleaned:
        return ExtractedDocument(
            status="no_text",
            page_count=pages,
            detail=(
                f"No text layer in this {pages}-page PDF — it is most likely a scan or an "
                f"image export. Nothing was read from it, so state the campaign details in "
                f"the chat instead."
            ),
        )

    return ExtractedDocument(
        status="ok",
        text=cleaned[:max_chars],
        char_count=min(len(cleaned), max_chars),
        page_count=pages,
        truncated=truncated or len(cleaned) > max_chars,
    )


# --------------------------------------------------------------- storage & retrieval
#
# Extraction happens once, at upload time, and the text lands in a sidecar file next to
# the staged document. Two reasons it is not re-parsed per agent turn: parsing a 50-page
# PDF on every tool call is wasted work against a rate limiter, and a rep needs to know
# *at upload* that their scanned PDF contributed nothing — not three minutes later when
# the package comes back without their constraints in it.
#
# The text goes to a file rather than into localDB because localDB is the app database and
# holds no document content. Only the summary is recorded there.

#: Suffix for the extracted-text sidecar.
TEXT_SUFFIX = ".extracted.txt"

_PREVIEW_CHARS = 200


def _stage_root() -> Path:
    """The directory `stored_path` values are relative to (the repo root)."""
    return get_settings().stage_dir.parent


def extract_and_store(stored: Path) -> DocumentSummary:
    """Parse an uploaded file and persist its text beside it.

    Never raises: a document that cannot be read still produces a summary saying so, and
    the upload itself stays successful. Losing the file because we could not parse it
    would be the worse outcome — the rep can still describe the brief in the chat.
    """
    result = extract_text(stored)
    summary = DocumentSummary(
        extraction_status=result.status,
        char_count=result.char_count,
        page_count=result.page_count,
        truncated=result.truncated,
        extraction_detail=result.detail,
    )
    if result.status != "ok" or not result.text:
        return summary

    sidecar = stored.with_name(stored.name + TEXT_SUFFIX)
    try:
        sidecar.write_text(result.text, encoding="utf-8")
        summary.text_path = str(sidecar.relative_to(_stage_root()))
    except OSError as exc:
        # The text was readable but could not be cached. Report it rather than claiming a
        # sidecar that is not there; the agent tool falls back to re-parsing.
        summary.extraction_detail = f"Text extracted but could not be cached ({exc.strerror})."

    preview = " ".join(result.text.split())[:_PREVIEW_CHARS]
    summary.preview = preview
    return summary


def load_text(record: dict) -> ExtractedDocument:
    """The extracted text for an upload record, from the sidecar or by re-parsing.

    Takes the localDB upload record so callers do not each reimplement path resolution.
    Falls back to parsing the original when the sidecar is missing — which happens for
    every document uploaded before extraction existed.
    """
    root = _stage_root()
    text_path = record.get("text_path")
    if text_path:
        candidate = root / str(text_path)
        if candidate.exists():
            text = candidate.read_text(encoding="utf-8", errors="replace")
            return ExtractedDocument(
                status="ok",
                text=text,
                char_count=len(text),
                page_count=record.get("page_count"),
                truncated=bool(record.get("truncated")),
            )

    stored_path = record.get("stored_path")
    if not stored_path:
        return ExtractedDocument(status="failed", detail="This upload has no stored path.")
    return extract_text(root / str(stored_path))
