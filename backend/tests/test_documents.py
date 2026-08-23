"""Brief parsing — .pdf, .docx, .txt.

The application accepted uploads for months and no parser ever read them: the prompt told
the agent to use "the filesystem tools", which in deepagents address a virtual state
filesystem, not the disk. So a rep could attach an RFP holding the budget, the flight dates
and the market list, and the package came back built from the one-line chat message.

These tests use real files — a hand-built PDF with a genuine xref table, a python-docx
document, encrypted and text-free PDFs — because the failure modes that matter here are all
in the file format, and a mocked parser would prove nothing about any of them.

Storage is redirected into a temp stage/ by the autouse fixture in conftest.py.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from app.config import get_settings
from app.main import app
from app.services import documents, local_db
from app.tools import master_tools


@pytest.fixture
def client() -> TestClient:
    return TestClient(app)


@pytest.fixture
def workdir(tmp_path: Path) -> Path:
    return tmp_path


# --------------------------------------------------------------------- fixtures


def write_pdf(path: Path, pages: list[str]) -> Path:
    """A minimal but *valid* PDF with a real xref table.

    Hand-built rather than pulled from a fixture binary or generated with reportlab: it
    keeps the test suite dependency-free and makes the page count explicit.
    """
    objects: list[bytes] = []
    kids = " ".join(f"{4 + 2 * i} 0 R" for i in range(len(pages)))
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for i, text in enumerate(pages):
        content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 3 0 R >> >> /Contents {5 + 2 * i} 0 R >>"
            ).encode()
        )
        objects.append(
            b"<< /Length "
            + str(len(content)).encode()
            + b" >>\nstream\n"
            + content
            + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode() + b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n"
    ).encode()

    path.write_bytes(bytes(out))
    return path


def write_docx(path: Path, paragraphs: list[str], table: list[list[str]] | None = None) -> Path:
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        built = document.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, value in enumerate(row):
                built.cell(r, c).text = value
    document.save(str(path))
    return path


# ------------------------------------------------------------------- extraction


def test_pdf_text_and_page_count(workdir: Path) -> None:
    path = write_pdf(
        workdir / "brief.pdf",
        ["Budget: USD 50,000 for 30 days", "Target: commuters aged 18-34"],
    )
    result = documents.extract_text(path)

    assert result.status == "ok"
    assert result.page_count == 2
    assert "50,000" in result.text
    assert "commuters aged 18-34" in result.text
    assert result.char_count == len(result.text)


def test_docx_reads_tables_not_only_paragraphs(workdir: Path) -> None:
    """A brief puts the budget in a table about as often as in a sentence.

    A paragraph-only reader drops precisely the fields intake needs, which is why this is
    pinned rather than left to the library's default traversal.
    """
    path = write_docx(
        workdir / "brief.docx",
        ["We want commuters in the Downtown Core zone."],
        table=[["Budget", "$50,000"], ["Start date", "2026-10-01"]],
    )
    result = documents.extract_text(path)

    assert result.status == "ok"
    assert "Downtown Core" in result.text
    assert "$50,000" in result.text
    assert "2026-10-01" in result.text


def test_txt_utf8(workdir: Path) -> None:
    path = workdir / "brief.txt"
    path.write_text("Budget: €20,000\nStart: 2026-10-01", encoding="utf-8")

    result = documents.extract_text(path)
    assert result.status == "ok"
    assert "€20,000" in result.text


def test_txt_cp1252_does_not_become_mojibake(workdir: Path) -> None:
    """A brief saved out of Word is cp1252, not utf-8, and it must not decode to garbage."""
    path = workdir / "brief.txt"
    path.write_bytes("Budget: £20,000 — “reach” focus".encode("cp1252"))

    result = documents.extract_text(path)
    assert result.status == "ok"
    assert "£20,000" in result.text
    assert "“reach”" in result.text
    assert "�" not in result.text


@pytest.mark.parametrize("suffix", [".xlsx", ".pptx", ".csv", ".md", ".exe", ""])
def test_unsupported_suffixes_are_reported_not_guessed(workdir: Path, suffix: str) -> None:
    path = workdir / f"file{suffix}"
    path.write_text("Budget: $50,000")

    result = documents.extract_text(path)
    assert result.status == "unsupported"
    assert result.text == ""


def test_missing_file_is_a_status_not_an_exception(workdir: Path) -> None:
    result = documents.extract_text(workdir / "gone.pdf")
    assert result.status == "failed"
    assert "no longer on disk" in result.detail


def test_empty_txt_is_no_text(workdir: Path) -> None:
    path = workdir / "empty.txt"
    path.write_text("   \n\n\t ")
    assert documents.extract_text(path).status == "no_text"


def test_pdf_with_no_text_layer_says_it_is_a_scan(workdir: Path) -> None:
    """The distinction a rep acts on: retype the numbers, versus resend the file."""
    path = write_pdf(workdir / "scan.pdf", [" "])
    result = documents.extract_text(path)

    assert result.status == "no_text"
    assert result.page_count == 1
    assert "scan" in result.detail.lower()


def test_password_protected_pdf_asks_for_an_unlocked_copy(workdir: Path) -> None:
    from pypdf import PdfWriter

    source = write_pdf(workdir / "plain.pdf", ["Budget: USD 50,000"])
    writer = PdfWriter(clone_from=str(source))
    writer.encrypt("a-user-password")
    locked = workdir / "locked.pdf"
    with locked.open("wb") as handle:
        writer.write(handle)

    result = documents.extract_text(locked)
    assert result.status == "failed"
    assert "password-protected" in result.detail


def test_owner_password_only_pdf_still_reads(workdir: Path) -> None:
    """Most "encrypted" PDFs in circulation open with an empty user password.

    Treating `is_encrypted` as unreadable would reject a large share of client decks that
    open fine in any viewer.
    """
    from pypdf import PdfWriter

    source = write_pdf(workdir / "plain.pdf", ["Budget: USD 50,000"])
    writer = PdfWriter(clone_from=str(source))
    writer.encrypt(user_password="", owner_password="owner-only")
    path = workdir / "owner.pdf"
    with path.open("wb") as handle:
        writer.write(handle)

    result = documents.extract_text(path)
    assert result.status == "ok"
    assert "50,000" in result.text


def test_corrupt_pdf_is_reported(workdir: Path) -> None:
    path = workdir / "corrupt.pdf"
    path.write_bytes(b"%PDF-1.4\nthis is not a pdf at all")
    assert documents.extract_text(path).status in {"failed", "no_text"}


def test_corrupt_docx_is_reported(workdir: Path) -> None:
    path = workdir / "corrupt.docx"
    path.write_bytes(b"not a zip archive")
    result = documents.extract_text(path)
    assert result.status == "failed"
    assert "corrupt" in result.detail


def test_cap_is_enforced_and_reported(workdir: Path) -> None:
    path = workdir / "long.txt"
    path.write_text("x" * 5000)

    result = documents.extract_text(path, max_chars=1000)
    assert result.status == "ok"
    assert len(result.text) <= 1000
    assert result.truncated is True


# ---------------------------------------------------------------------- cleaning


@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        # Zero-width and BOM carry no text.
        ("a​b﻿c", "abc"),
        # A non-breaking space is a space; deleting it would join two words.
        ("Downtown Core", "Downtown Core"),
        # U+2028 *is* a line break — dropping it ran a budget onto a date.
        ("Budget: 50k Start: Oct 1", "Budget: 50k\nStart: Oct 1"),
        # Control characters break JSON transport; tab and newline survive.
        ("x\x00\x07y\tz", "xy\tz"),
        # PDF page breaks produce runs of blank lines, which are pure token cost.
        ("a\n\n\n\n\nb", "a\n\nb"),
        ("trailing   \nspace\t\n", "trailing\nspace"),
        ("\r\nwindows\r\nnewlines\r\n", "windows\nnewlines"),
    ],
)
def test_clean_text(raw: str, expected: str) -> None:
    assert documents.clean_text(raw) == expected


# ----------------------------------------------------------------------- excerpt


def test_excerpt_returns_short_text_whole() -> None:
    assert documents.excerpt("a short brief", limit=100) == ("a short brief", False)


def test_excerpt_cuts_on_a_paragraph_boundary() -> None:
    """A mid-sentence cut invites the model to complete the thought.

    Completing "the budget is $5" into a number is the single worst thing this system can
    do, so the cut lands on a boundary.
    """
    body = "This is the opening section of the brief and it runs on for a while."
    text = body + "\n\n" + "second paragraph " * 40
    out, truncated = documents.excerpt(text, limit=90)

    assert truncated is True
    assert out == body


def test_excerpt_falls_back_to_a_word_boundary() -> None:
    """No line break nearby still must not split a token — a cut number is a fabrication."""
    text = "the budget is $50,000 and " * 40
    out, truncated = documents.excerpt(text, limit=100)

    assert truncated is True
    assert not out.endswith("$50,0")
    # Whole tokens only: the excerpt is a clean prefix of the document.
    assert text.startswith(out)


def test_excerpt_hard_cuts_text_with_no_whitespace() -> None:
    """A 500-character unbroken string has no boundary to honour."""
    out, truncated = documents.excerpt("x" * 500, limit=100)
    assert truncated is True
    assert len(out) == 100


def test_excerpt_ignores_an_early_boundary() -> None:
    text = "Hi\n\n" + "y" * 500
    out, _ = documents.excerpt(text, limit=100)
    # Cutting at the boundary would return 2 characters of a 500-character document.
    assert len(out) > 50


# ------------------------------------------------------------ store & load back


def test_extract_and_store_writes_a_sidecar(workdir: Path) -> None:
    staged = get_settings().stage_dir / "ses-test"
    staged.mkdir(parents=True, exist_ok=True)
    path = write_pdf(staged / "brief.pdf", ["Budget: USD 50,000"])

    summary = documents.extract_and_store(path)

    assert summary.extraction_status == "ok"
    assert summary.page_count == 1
    assert summary.text_path is not None
    assert summary.preview.startswith("Budget")

    sidecar = get_settings().stage_dir.parent / summary.text_path
    assert sidecar.exists()
    assert "50,000" in sidecar.read_text(encoding="utf-8")


def test_unreadable_file_stores_no_sidecar(workdir: Path) -> None:
    staged = get_settings().stage_dir / "ses-test"
    staged.mkdir(parents=True, exist_ok=True)
    path = write_pdf(staged / "scan.pdf", [" "])

    summary = documents.extract_and_store(path)
    assert summary.extraction_status == "no_text"
    assert summary.text_path is None
    assert summary.preview == ""


def test_load_text_prefers_the_sidecar(workdir: Path) -> None:
    """The sidecar is the cache; re-parsing per turn would be wasted work."""
    staged = get_settings().stage_dir / "ses-test"
    staged.mkdir(parents=True, exist_ok=True)
    path = write_pdf(staged / "cached.pdf", ["Original text"])
    summary = documents.extract_and_store(path)

    sidecar = get_settings().stage_dir.parent / summary.text_path
    sidecar.write_text("SIDECAR WINS", encoding="utf-8")

    record = {
        "stored_path": str(path.relative_to(get_settings().stage_dir.parent)),
        **summary.model_dump(mode="json"),
    }
    assert documents.load_text(record).text == "SIDECAR WINS"


def test_load_text_reparses_when_the_sidecar_is_gone(workdir: Path) -> None:
    """Every document uploaded before extraction existed takes this path."""
    staged = get_settings().stage_dir / "ses-test"
    staged.mkdir(parents=True, exist_ok=True)
    path = write_pdf(staged / "legacy.pdf", ["Budget: USD 12,345"])

    record = {"stored_path": str(path.relative_to(get_settings().stage_dir.parent))}
    result = documents.load_text(record)

    assert result.status == "ok"
    assert "12,345" in result.text


def test_load_text_without_a_path_is_a_status() -> None:
    assert documents.load_text({}).status == "failed"


# --------------------------------------------------------------------- endpoints


def _upload(client: TestClient, session_id: str, name: str, data: bytes) -> dict:
    return client.post(
        "/uploads",
        data={"session_id": session_id},
        files={"file": (name, data, "application/octet-stream")},
    ).json()


def test_upload_returns_the_extraction_result(client: TestClient, workdir: Path) -> None:
    session_id = local_db.insert(local_db.SESSIONS, {"title": "t"})["id"]
    pdf = write_pdf(workdir / "brief.pdf", ["Budget: USD 50,000 over 30 days"])

    body = _upload(client, session_id, "brief.pdf", pdf.read_bytes())

    assert body["extraction_status"] == "ok"
    assert body["page_count"] == 1
    assert body["char_count"] > 0
    assert "Budget" in body["preview"]


def test_upload_of_a_scan_reports_it_at_upload_time(client: TestClient, workdir: Path) -> None:
    session_id = local_db.insert(local_db.SESSIONS, {"title": "t"})["id"]
    pdf = write_pdf(workdir / "scan.pdf", [" "])

    body = _upload(client, session_id, "scan.pdf", pdf.read_bytes())
    assert body["extraction_status"] == "no_text"
    assert "scan" in body["extraction_detail"].lower()


@pytest.mark.parametrize("name", ["sheet.xlsx", "deck.pptx", "data.csv", "notes.md"])
def test_formats_with_no_parser_are_rejected(client: TestClient, name: str) -> None:
    """Accepting a file nothing can read is worse than refusing it.

    The upload used to succeed, the agent saw nothing, and no one was told.
    """
    session_id = local_db.insert(local_db.SESSIONS, {"title": "t"})["id"]
    response = client.post(
        "/uploads",
        data={"session_id": session_id},
        files={"file": (name, b"anything", "application/octet-stream")},
    )
    assert response.status_code == 415


def test_delete_upload_removes_the_bytes_and_the_text(client: TestClient, workdir: Path) -> None:
    session_id = local_db.insert(local_db.SESSIONS, {"title": "t"})["id"]
    pdf = write_pdf(workdir / "brief.pdf", ["Budget: USD 50,000"])
    body = _upload(client, session_id, "brief.pdf", pdf.read_bytes())

    record = local_db.get_record(local_db.UPLOADS, body["id"])
    root = get_settings().stage_dir.parent
    stored = root / record["stored_path"]
    sidecar = root / record["text_path"]
    assert stored.exists() and sidecar.exists()

    assert client.delete(f"/uploads/{body['id']}").status_code == 200
    assert not stored.exists()
    assert not sidecar.exists()
    assert local_db.get_record(local_db.UPLOADS, body["id"]) is None


def test_get_unknown_upload_is_404(client: TestClient) -> None:
    assert client.get("/uploads/upl-nope").status_code == 404


# -------------------------------------------------------- the agent-facing tool


def _stage(session_id: str, name: str, pages: list[str], workdir: Path) -> dict:
    staged = get_settings().stage_dir / session_id
    staged.mkdir(parents=True, exist_ok=True)
    path = write_pdf(staged / name, pages)
    summary = documents.extract_and_store(path)
    return local_db.insert(
        local_db.UPLOADS,
        {
            "session_id": session_id,
            "filename": name,
            "size_bytes": path.stat().st_size,
            "stored_path": str(path.relative_to(get_settings().stage_dir.parent)),
            **summary.model_dump(mode="json"),
        },
    )


def test_tool_returns_the_document_text(workdir: Path) -> None:
    record = _stage("ses-tool", "brief.pdf", ["Budget: USD 50,000 for 30 days"], workdir)
    out = master_tools.read_campaign_document.invoke({"upload_id": record["id"]})

    assert out["status"] == "ok"
    assert out["filename"] == "brief.pdf"
    assert "50,000" in out["text"]
    assert out["truncated"] is False
    assert out["returned_chars"] <= out["total_chars"]


def test_tool_refuses_to_guess_an_id() -> None:
    out = master_tools.read_campaign_document.invoke({"upload_id": "upl-nope"})
    assert out["status"] == "not_found"
    assert "do not guess" in out["detail"]


def test_tool_short_circuits_a_known_unreadable_document(workdir: Path) -> None:
    """A recorded failure will not change on a re-read, so it must not cost the I/O."""
    record = _stage("ses-tool", "scan.pdf", [" "], workdir)
    out = master_tools.read_campaign_document.invoke({"upload_id": record["id"]})

    assert out["status"] == "no_text"
    assert "text" not in out
    # The instruction that stops the model inventing contents from the filename.
    assert "do not" in out["guidance"].lower()


def test_tool_never_returns_more_than_the_excerpt_cap(workdir: Path) -> None:
    """SOLUTION.md §31: a whole document never enters agent context."""
    session_id = "ses-big"
    staged = get_settings().stage_dir / session_id
    staged.mkdir(parents=True, exist_ok=True)
    path = staged / "long.txt"
    path.write_text("Budget line.\n" * 5000, encoding="utf-8")

    summary = documents.extract_and_store(path)
    record = local_db.insert(
        local_db.UPLOADS,
        {
            "session_id": session_id,
            "filename": "long.txt",
            "size_bytes": path.stat().st_size,
            "stored_path": str(path.relative_to(get_settings().stage_dir.parent)),
            **summary.model_dump(mode="json"),
        },
    )

    out = master_tools.read_campaign_document.invoke({"upload_id": record["id"]})
    assert out["status"] == "ok"
    assert len(out["text"]) <= documents.AGENT_EXCERPT_CHARS
    assert out["truncated"] is True


def test_the_tool_is_registered() -> None:
    assert "read_campaign_document" in [t.name for t in master_tools.TOOLS]
